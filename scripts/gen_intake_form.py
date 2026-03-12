#!/usr/bin/env python3
"""
Generate Avanti brand intake form as a styled Word .docx
Output: docs/Avanti-品牌信息收集表.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

ORANGE  = RGBColor(0xFF, 0x6B, 0x35)
NAVY    = RGBColor(0x0D, 0x1B, 0x2E)
MUTED   = RGBColor(0x70, 0x70, 0xA0)
LIGHT_BG = "F4F5F8"
HDR_BG   = "0D1B2E"


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_padding(cell, top=80, bottom=80, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def add_section_header(doc, number: str, title: str):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    # Number badge
    r1 = p.add_run(f"  {number}  ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Hack: use highlight won't work for custom color; use separate table row for header below
    # Title
    r2 = p.add_run(f"  {title}")
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = ORANGE
    return p


def create_table(doc, col_widths=(Inches(2.1), Inches(4.4))):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for i, w in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = w
    return table


def add_row(table, label: str, placeholder: str = "", height_pt: int = 22):
    row = table.add_row()

    lc = row.cells[0]
    set_cell_bg(lc, LIGHT_BG)
    set_cell_padding(lc)
    lp = lc.paragraphs[0]
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(9.5)

    vc = row.cells[1]
    set_cell_padding(vc)
    if placeholder:
        vp = vc.paragraphs[0]
        vr = vp.add_run(placeholder)
        vr.font.size = Pt(9)
        vr.font.color.rgb = RGBColor(0xBB, 0xBB, 0xCC)

    return row


def add_checkbox_row(table, label: str, options: list[str]):
    row = table.add_row()

    lc = row.cells[0]
    set_cell_bg(lc, LIGHT_BG)
    set_cell_padding(lc)
    lp = lc.paragraphs[0]
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(9.5)

    vc = row.cells[1]
    set_cell_padding(vc, top=60, bottom=60)
    for i, opt in enumerate(options):
        if i == 0:
            vp = vc.paragraphs[0]
        else:
            vp = vc.add_paragraph()
        vp.paragraph_format.space_before = Pt(1)
        vp.paragraph_format.space_after = Pt(1)
        vr = vp.add_run(f"☐  {opt}")
        vr.font.size = Pt(9)

    return row


def build(output_path: Path):
    doc = Document()

    # ── Page setup ──────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width  = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = sec.right_margin = Inches(0.9)
    sec.top_margin  = sec.bottom_margin = Inches(0.75)

    # Default body font
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    # ── Title block ─────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t.paragraph_format.space_after = Pt(2)
    r = t.add_run("AVANTI")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = ORANGE

    sub = doc.add_paragraph()
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after = Pt(2)
    sr = sub.add_run("AI 可见度评估 · 品牌信息收集表")
    sr.font.size = Pt(13)
    sr.bold = True
    sr.font.color.rgb = NAVY

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(6)
    rr = rule.add_run("─" * 72)
    rr.font.size = Pt(7)
    rr.font.color.rgb = MUTED

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(12)
    ir = intro.add_run(
        "请填写以下信息，我们将为您生成第一份完整的 AI 品牌基准诊断报告，之后每周持续追踪。\n"
        "报告包含：GEO Score · 四大 AI 引擎表现 · 竞品对比 · 幻觉检测 · 选品情报 · 优化行动计划"
    )
    ir.font.size = Pt(9.5)
    ir.font.color.rgb = MUTED

    # ── PART 1: Brand basics ────────────────────────────
    add_section_header(doc, "PART 1", "品牌基础信息")
    t1 = create_table(doc)
    add_row(t1, "品牌英文名称 *",      "商标注册英文名，如：Jackery")
    add_row(t1, "品牌中文名称",         "如有，如：杰克瑞")
    add_row(t1, "官网 / 独立站链接",          "https://（没有填 N/A）")
    add_row(t1, "TikTok Shop 店铺链接 *",    "https://")
    add_row(t1, "Shopee 店铺链接",           "https://（没有填 N/A）")
    add_checkbox_row(t1, "主要销售平台 *", [
        "TikTok Shop", "Shopee", "独立站", "其他：___________"
    ])
    add_checkbox_row(t1, "主要目标市场 *", [
        "北美（美/加）", "欧洲", "日本/韩国", "东南亚 — 泰国",
        "东南亚 — 马来西亚", "东南亚 — 菲律宾", "其他：___________"
    ])

    # ── PART 2: Products ────────────────────────────────
    add_section_header(doc, "PART 2", "产品信息")
    t2 = create_table(doc)
    add_row(t2, "主营品类 *",            "如：便携储能 / 蓝牙耳机 / 户外装备 / 美妆")
    add_row(t2, "核心产品英文名 *（最多 5 个）",
            "有几个填几个\n1.\n2.\n3.\n4.\n5.")
    add_row(t2, "产品页面链接（TikTok / Shopee）*",
            "主力产品的链接即可，粘 1-3 条\n1.\n2.\n3.")
    add_row(t2, "产品核心卖点 / 差异化",  "3-5 个关键词，如：2000W大功率 · 双向快充 · IP67防水")
    add_row(t2, "近期合作达人内容链接（可选）",
            "粘 1-3 条视频 / 测评链接，我们会检测这些内容有没有被 AI 引用\n1.\n2.\n3.")
    add_checkbox_row(t2, "月均 GMV 区间 *", [
        "< $2 万", "$2 万 – $10 万", "$10 万 – $50 万", "> $50 万"
    ])

    # ── PART 3: Competitors ─────────────────────────────
    add_section_header(doc, "PART 3", "竞争对手")
    t3 = create_table(doc)
    add_row(t3, "主要竞品品牌 Top 3 *",
            "英文品牌名\n1.\n2.\n3.")
    add_row(t3, "最担心被哪个品牌超越",  "品牌名 + 一句话原因")

    # ── PART 4: Current AI awareness ────────────────────
    add_section_header(doc, "PART 4", "当前 AI 感知状况")
    t4 = create_table(doc)
    add_checkbox_row(t4, "是否问过 AI 推荐你们品类的产品？", [
        "经常（每周）", "偶尔", "从来没有"
    ])
    add_checkbox_row(t4, "AI 推荐过你们品牌吗？", [
        "是，经常出现", "偶尔出现", "从不出现", "不确定"
    ])
    add_checkbox_row(t4, "AI 是否说过关于你们产品的错误信息？", [
        "有（请在旁边备注）：___________", "没有", "从未查过"
    ])
    add_checkbox_row(t4, "目前用什么工具监控 AI 表现？", [
        "没有任何工具", "TikTok 后台数据", "SEMrush / Ahrefs", "其他：___________"
    ])

    # ── PART 5: Goals ───────────────────────────────────
    add_section_header(doc, "PART 5", "这次合作，你最想了解什么？（可多选）")
    t5 = create_table(doc)
    add_checkbox_row(t5, "优先关注项", [
        "我的品牌 GEO Score 和 AI 引擎排名",
        "竞品在 AI 里的表现和差距",
        "AI 有没有说错关于我们产品的信息（幻觉检测）",
        "当前 AI 正在向买家推荐哪些品类 / 产品（选品情报）",
        "如何提升 AI 可见度的具体可执行建议",
    ])
    add_row(t5, "其他补充 / 问题",
            "任何你觉得我们需要知道的信息，或者想问的问题")

    # ── PART 6: Contact ─────────────────────────────────
    add_section_header(doc, "PART 6", "联系人信息")
    t6 = create_table(doc)
    add_row(t6, "姓名 *",           "")
    add_row(t6, "职位 / 角色",       "如：品牌总监 / 运营负责人 / CEO")
    add_row(t6, "微信号 / 手机 *",   "")
    add_row(t6, "邮箱",              "用于接收报告")

    # ── Footer ──────────────────────────────────────────
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Avanti  ·  avantia2a.com  ·  hello@avantia2a.com  ·  © 2026 Avanti Growth Labs LLC")
    fr.font.size = Pt(8)
    fr.font.color.rgb = MUTED
    fr.italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"✓ Saved: {output_path}")


if __name__ == "__main__":
    out = Path("/Users/johnsonliu/Desktop/claude code/ai-rec-monitor/docs/Avanti-品牌信息收集表.docx")
    build(out)
